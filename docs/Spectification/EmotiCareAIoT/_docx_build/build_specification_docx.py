from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
REFERENCE = Path(__file__).resolve().parent / "specification-cover-reference.docx"
SOURCE = ROOT / "Specification.md"
CHARTS = Path(__file__).resolve().parent / "charts"
OUTPUT = ROOT / "specification.docx"

FONT = "Times New Roman"
ACCENT = "1F4E79"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    tr_pr.append(el)


def apply_run(run, size=13, bold=None, italic=None, color=None, code=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if code:
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")


INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")


def add_inline(paragraph, text: str, size=13, color=None):
    parts = INLINE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            apply_run(run, size, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            apply_run(run, max(10, size - 1), code=True, color=color)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            apply_run(run, size, italic=True, color=color)
        else:
            run = paragraph.add_run(part)
            apply_run(run, size, color=color)


def format_paragraph(paragraph, before=0, after=6, line=1.15, left=0, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.left_indent = Cm(left) if left else None
    fmt.keep_with_next = keep


def add_heading(doc, text: str, level: int):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    if level == 1:
        p.paragraph_format.page_break_before = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_inline(p, text, 21, ACCENT)
        for run in p.runs:
            run.bold = True
        format_paragraph(p, before=0, after=12, line=1.08, keep=True)
    elif level == 2:
        add_inline(p, text, 18, ACCENT)
        for run in p.runs:
            run.bold = True
        format_paragraph(p, before=14, after=8, line=1.08, keep=True)
    elif level == 3:
        add_inline(p, text, 15, ACCENT)
        for run in p.runs:
            run.bold = True
        format_paragraph(p, before=11, after=6, line=1.1, keep=True)
    else:
        add_inline(p, text, 13, ACCENT)
        for run in p.runs:
            run.bold = True
        format_paragraph(p, before=8, after=4, line=1.1, keep=True)
    return p


def split_table_row(line: str):
    return [x.strip() for x in line.strip().strip("|").split("|")]


def is_separator(row):
    return all(re.fullmatch(r":?-{3,}:?", x.replace(" ", "")) for x in row)


def add_table(doc, rows):
    headers = split_table_row(rows[0])
    data = [split_table_row(x) for x in rows[2:] if x.strip()]
    col_count = len(headers)
    data = [r + [""] * (col_count - len(r)) for r in data]
    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for j, value in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "D9EAF7")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_paragraph(p, after=0, line=1.0)
        add_inline(p, value, 10.5)
        for r in p.runs:
            r.bold = True
    set_repeat_table_header(table.rows[0])
    for row_data in data:
        cells = table.add_row().cells
        for j, value in enumerate(row_data):
            cell = cells[j]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            format_paragraph(p, after=0, line=1.0)
            add_inline(p, value, 10.5)
    p = doc.add_paragraph()
    format_paragraph(p, after=4)


def ensure_numbering(doc, ordered: bool):
    numbering = doc.part.numbering_part.element
    num_id = "910" if ordered else "911"
    if numbering.xpath(f'./w:num[@w:numId="{num_id}"]'):
        return num_id
    abstract_id = "910" if ordered else "911"
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), abstract_id)
    for level in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
        lvl.append(fmt)
        txt = OxmlElement("w:lvlText")
        txt.set(qn("w:val"), f"%{level + 1}." if ordered else "•")
        lvl.append(txt)
        ppr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(720 + level * 360))
        ind.set(qn("w:hanging"), "360")
        ppr.append(ind)
        lvl.append(ppr)
        abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list(doc, text, ordered=False, level=0):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(min(level, 2)))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), ensure_numbering(doc, ordered))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    add_inline(p, text, 13)
    format_paragraph(p, after=2, line=1.1)


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("\n".join(lines))
    apply_run(r, 9.5, code=True)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    p_pr.append(shd)


def add_figure(doc, chart_no):
    image = CHARTS / f"chart-{chart_no:02d}.png"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(image), width=Cm(16.5))
    format_paragraph(p, before=4, after=2, keep=True)
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(c, f"Sơ đồ {chart_no}. Minh họa luồng xử lý tương ứng.", 11, "666666")
    for run in c.runs:
        run.italic = True
    format_paragraph(c, after=8, line=1.0)


def build():
    doc = Document(REFERENCE)
    section = doc.sections[0]
    # The cover occupies the retained template; the following body is a new page.
    doc.add_page_break()
    chart_no = 0
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        if not stripped or stripped == "---":
            i += 1
            continue
        if stripped.startswith("```"):
            language = stripped[3:].strip().lower()
            block = []
            i += 1
            while i < len(lines) and lines[i].strip() != "```":
                block.append(lines[i])
                i += 1
            if language == "mermaid":
                chart_no += 1
                add_figure(doc, chart_no)
            else:
                add_code_block(doc, block)
            i += 1
            continue
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            add_heading(doc, m.group(2), len(m.group(1)))
            i += 1
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                rows.append(lines[i].strip())
                i += 1
            if len(rows) >= 2 and is_separator(split_table_row(rows[1])):
                add_table(doc, rows)
            else:
                for row in rows:
                    p = doc.add_paragraph()
                    add_inline(p, row, 13)
                    format_paragraph(p)
            continue
        m = re.match(r"^(\s*)[-*+]\s+(.*)$", raw)
        if m:
            add_list(doc, m.group(2), False, len(m.group(1)) // 2)
            i += 1
            continue
        m = re.match(r"^(\s*)\d+[.)]\s+(.*)$", raw)
        if m:
            add_list(doc, m.group(2), True, len(m.group(1)) // 2)
            i += 1
            continue
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            add_inline(p, stripped[2:], 13)
            for run in p.runs:
                run.italic = True
            format_paragraph(p, before=3, after=6, left=0.7)
            i += 1
            continue
        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            look = lines[i].strip()
            if (not look or look == "---" or look.startswith("```") or re.match(r"^(#{1,4})\s+", look)
                    or (look.startswith("|") and look.endswith("|"))
                    or re.match(r"^\s*[-*+]\s+", lines[i]) or re.match(r"^\s*\d+[.)]\s+", lines[i])
                    or look.startswith("> ")):
                break
            paragraph_lines.append(look)
            i += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(paragraph_lines), 13)
        format_paragraph(p, after=6, line=1.15)
    if chart_no != 8:
        raise RuntimeError(f"Expected 8 charts, found {chart_no}")
    doc.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    build()
