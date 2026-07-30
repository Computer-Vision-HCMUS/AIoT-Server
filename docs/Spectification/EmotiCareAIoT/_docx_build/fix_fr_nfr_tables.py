from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent.parent
SOURCE_DOCX = ROOT / "specification.docx"
SOURCE_MD = ROOT / "Specification.md"
OUTPUT = ROOT / "sup-spec1.docx"
FONT = "Times New Roman"


def markdown_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator(row: list[str]) -> bool:
    return bool(row) and all(re.fullmatch(r":?-{2,}:?", cell.replace(" ", "")) for cell in row)


def get_requirement_tables() -> dict[str, list[list[str]]]:
    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    result: dict[str, list[list[str]]] = {}
    i = 0
    while i < len(lines):
        match = re.match(r"^##\s+(5\.[2-8]\.|6\.[2-7]\.)\s+(.*)$", lines[i])
        if not match:
            i += 1
            continue
        heading = f"{match.group(1)} {match.group(2)}"
        j = i + 1
        while j < len(lines) and not lines[j].strip():
            j += 1
        rows = []
        while j < len(lines) and lines[j].strip().startswith("|") and lines[j].strip().endswith("|"):
            rows.append(markdown_row(lines[j]))
            j += 1
        if len(rows) >= 3 and is_separator(rows[1]):
            result[heading] = [rows[0]] + rows[2:]
        i = j
    if len(result) != 13:
        raise RuntimeError(f"Expected 13 FR/NFR tables, found {len(result)}")
    return result


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = tc_mar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_run(run, size: float, bold=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold


def populate_cell(cell, text: str, size: float, bold=False, center=False):
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0
    run = p.add_run(text.replace("`", ""))
    set_run(run, size, bold)


def set_table_layout(table, widths: list[float]):
    available = 9036  # A4, 1-inch left/right margins
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(available))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_width(cell, int(available * widths[idx]))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def make_table(doc, rows: list[list[str]], widths: list[float], body_size=9.5):
    col_count = len(rows[0])
    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"
    for col, value in enumerate(rows[0]):
        cell = table.rows[0].cells[col]
        set_cell_shading(cell, "D9EAF7")
        populate_cell(cell, value, body_size, bold=True, center=True)
    set_repeat_header(table.rows[0])
    for row_values in rows[1:]:
        cells = table.add_row().cells
        values = row_values + [""] * (col_count - len(row_values))
        for col, value in enumerate(values[:col_count]):
            populate_cell(cells[col], value, body_size, center=(col in (0, col_count - 1)))
    set_table_layout(table, widths)
    return table


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def replace_literal_table(doc, heading_text: str, rows: list[list[str]]):
    heading = next((p for p in doc.paragraphs if p.text.strip() == heading_text), None)
    if heading is None:
        raise RuntimeError(f"Heading not found: {heading_text}")
    current = heading._p.getnext()
    to_remove = []
    while current is not None and current.tag == qn("w:p"):
        from docx.text.paragraph import Paragraph
        p = Paragraph(current, heading._parent)
        if not p.text.strip().startswith("|"):
            break
        to_remove.append(p)
        current = current.getnext()
    if len(to_remove) < 3:
        raise RuntimeError(f"Literal table not found after: {heading_text}")
    for p in to_remove:
        remove_paragraph(p)
    widths = [0.09, 0.63, 0.17, 0.11]
    table = make_table(doc, rows, widths)
    heading._p.addnext(table._tbl)


def insert_after_paragraph(doc, heading_text: str, rows: list[list[str]]):
    heading = next((p for p in doc.paragraphs if p.text.strip() == heading_text), None)
    if heading is None:
        raise RuntimeError(f"Heading not found: {heading_text}")
    table = make_table(doc, rows, [0.08, 0.19, 0.19, 0.18, 0.20, 0.16], body_size=8.4)
    heading._p.addnext(table._tbl)


def main():
    shutil.copy2(SOURCE_DOCX, OUTPUT)
    doc = Document(OUTPUT)
    tables = get_requirement_tables()
    for heading, rows in tables.items():
        replace_literal_table(doc, heading, rows)
    reference_rows = [
        ["ID", "Tình huống sử dụng", "Đầu vào", "Đầu ra", "Xử lý chính", "Mục tiêu thời gian"],
        ["UC-01", "Nhận diện cảm xúc bằng giọng nói", "Giọng nói người dùng", "Nhãn cảm xúc và độ tin cậy", "Xử lý tại thiết bị", "Không quá 30 giây"],
        ["UC-02", "Gợi ý hoạt động cải thiện tâm trạng", "Kết quả cảm xúc và lịch sử đã đồng bộ", "5 thẻ hoạt động trên màn hình", "Máy chủ và màn hình thiết bị", "Không quá 20 giây khi có Internet"],
        ["UC-03", "Lựa chọn bài hát hoặc podcast theo chủ đích", "Nhóm nội dung, loại nội dung, chủ đích và ngữ cảnh cảm xúc nếu có", "Danh sách bài hát hoặc podcast trên màn hình", "Máy chủ và màn hình thiết bị", "Không quá 20 giây khi có Internet"],
        ["UC-04", "Trò chuyện hỗ trợ cảm xúc", "Giọng nói hoặc câu hỏi và ngữ cảnh cảm xúc nếu có", "Thẻ phản hồi trên màn hình", "Máy chủ và màn hình thiết bị", "Không quá 20 giây khi có Internet"],
        ["UC-05", "Thống kê và phân tích xu hướng cảm xúc", "Lịch sử cảm xúc, hoạt động, nội dung đã chọn và thông tin trò chuyện", "Bản tóm tắt trên màn hình", "Máy chủ và màn hình thiết bị", "Không quá 180 giây"],
    ]
    insert_after_paragraph(doc, "9.2. Bảng tham chiếu tình huống sử dụng", reference_rows)
    doc.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    main()
